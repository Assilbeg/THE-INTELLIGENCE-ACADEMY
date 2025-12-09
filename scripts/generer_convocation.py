#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur de convocations pour formations MINDNESS
Utilise le template "convocation template.docx" comme base.
Génère une convocation par apprenant.
"""

import json
import sys
import os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.table import _Cell
from docx.shared import Cm, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from copy import deepcopy


def format_date_fr(date_obj):
    """Formate une date en français (ex: 16 décembre 2024)."""
    mois = [
        "", "janvier", "février", "mars", "avril", "mai", "juin",
        "juillet", "août", "septembre", "octobre", "novembre", "décembre"
    ]
    return f"{date_obj.day} {mois[date_obj.month]} {date_obj.year}"


def parse_date(date_str):
    """Parse une date en différents formats."""
    formats = ["%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%d/%m/%y"]
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    raise ValueError(f"Format de date non reconnu: {date_str}")


def replace_in_paragraph(para, old_text, new_text):
    """Remplace le texte dans un paragraphe en préservant le formatage si possible."""
    if old_text not in para.text:
        return False
    
    # Essayer de remplacer dans chaque run
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    
    # Si fragmenté entre plusieurs runs, reconstruire le paragraphe
    new_full_text = para.text.replace(old_text, new_text)
    para.clear()
    para.add_run(new_full_text)
    return True


def copy_row(table, row_idx):
    """Copie une ligne de tableau et l'ajoute à la fin."""
    tbl = table._tbl
    tr = table.rows[row_idx]._tr
    new_tr = deepcopy(tr)
    tbl.append(new_tr)
    return table.rows[-1]


def generer_convocation(data: dict, output_dir: str = None, template_path: str = "convocation template.docx"):
    """
    Génère une convocation par apprenant.
    """
    
    # Trouver le template
    template = Path(template_path)
    if not template.exists():
        script_dir = Path(__file__).parent
        template = script_dir.parent / "templates" / template_path
        if not template.exists():
            raise FileNotFoundError(f"Template non trouvé: {template_path}")
    template_path = str(template)
    
    # Parser les dates
    date_debut = parse_date(data["date_debut"])
    date_fin = parse_date(data["date_fin"])
    
    # Date d'émission : personnalisée ou aujourd'hui par défaut
    if "date_emission" in data and data["date_emission"]:
        date_emission = parse_date(data["date_emission"])
    else:
        date_emission = datetime.now()
    
    # Lien ressources (optionnel)
    lien_ressources = data.get("lien_ressources", "")
    
    # Formateur(s)
    formateurs = data.get("formateurs", [])
    if isinstance(formateurs, list):
        formateur_str = ", ".join(formateurs)
    else:
        formateur_str = formateurs
    
    # Sessions pour le tableau
    sessions = data.get("sessions", [])
    
    # Lieu (pour le tableau si pas spécifié dans les sessions)
    lieu_general = data.get("lieu", "")
    
    print(f"📧 Génération des convocations")
    print(f"   Formation : {data['nom_formation']}")
    print(f"   Du {format_date_fr(date_debut)} au {format_date_fr(date_fin)}")
    print(f"   Durée : {data['duree_heures']} heures ({data['duree_jours']} jours)")
    print(f"   Lieu : {lieu_general}")
    print(f"   Sessions : {len(sessions)}")
    print(f"   {len(data['apprenants'])} apprenant(s)")
    
    fichiers_generes = []
    
    for apprenant in data["apprenants"]:
        nom = apprenant["nom"]
        prenom = apprenant["prenom"]
        
        # Charger le template
        doc = Document(template_path)
        
        # Préparer les remplacements de base
        replacements = {
            "{{DATE_EMISSION}}": format_date_fr(date_emission),
            "{{PRENOM}}": prenom,
            "{{NOM}}": nom,
            "{{NOM_FORMATION}}": data["nom_formation"],
            "{{LIEU}}": lieu_general,
            "{{DATE_DEBUT}}": format_date_fr(date_debut),
            "{{DATE_FIN}}": format_date_fr(date_fin),
            "{{DUREE_HEURES}}": str(data["duree_heures"]),
            "{{DUREE_JOURS}}": str(data["duree_jours"]),
            "{{FORMATEUR}}": formateur_str,
        }
        
        # Remplacements dans les paragraphes
        for para in doc.paragraphs:
            # Traitement spécial pour le paragraphe Lieu/Dates/Durée
            if "{{LIEU}}" in para.text and "{{DATE_DEBUT}}" in para.text:
                para.clear()
                # Lieu
                para.add_run("Lieu de la formation : ").font.name = "Calibri"
                run_lieu = para.add_run(f"{lieu_general}.")
                run_lieu.font.name = "Calibri"
                run_lieu.font.bold = True
                # Dates
                para.add_run("\nDates de la formation : du ").font.name = "Calibri"
                run_dates = para.add_run(f"{format_date_fr(date_debut)} au {format_date_fr(date_fin)}.")
                run_dates.font.name = "Calibri"
                run_dates.font.bold = True
                # Durée
                para.add_run("\nDurée de la formation : ").font.name = "Calibri"
                run_duree = para.add_run(f"{data['duree_heures']} heures ({data['duree_jours']} jours).")
                run_duree.font.name = "Calibri"
                run_duree.font.bold = True
                continue
            
            # Traitement spécial pour le formateur
            if "{{FORMATEUR}}" in para.text:
                para.clear()
                para.add_run("Formateur(trice) : ").font.name = "Calibri"
                run_formateur = para.add_run(formateur_str)
                run_formateur.font.name = "Calibri"
                run_formateur.font.bold = True
                continue
            
            # Autres remplacements standards
            for old, new in replacements.items():
                replace_in_paragraph(para, old, new)
            
            # Gérer le lien ressources (optionnel)
            if "{{LIEN_RESSOURCES}}" in para.text:
                if lien_ressources:
                    replace_in_paragraph(para, "{{LIEN_RESSOURCES}}", lien_ressources)
                else:
                    para.clear()
        
        # Remplacements dans le tableau 0 (nom de formation) - en gras
        if len(doc.tables) > 0:
            for row in doc.tables[0].rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "{{NOM_FORMATION}}" in para.text:
                            para.clear()
                            run = para.add_run(data["nom_formation"])
                            run.font.name = "Calibri"
                            run.font.bold = True
        
        # Générer le tableau 1 (planning) dynamiquement
        if len(doc.tables) > 1 and sessions:
            table = doc.tables[1]
            
            # Supprimer les lignes de données existantes (garder l'en-tête ligne 0)
            while len(table.rows) > 1:
                tr = table.rows[-1]._tr
                table._tbl.remove(tr)
            
            # Définir les largeurs des colonnes (Date: 5cm, Heure: 6cm, Lieu: 5cm)
            table.columns[0].width = Cm(5)
            table.columns[1].width = Cm(6)
            table.columns[2].width = Cm(5)
            
            # Récupérer le style de l'en-tête pour l'appliquer aux nouvelles lignes
            header_row = table.rows[0]
            
            # Ajouter une ligne par session
            for session in sessions:
                # Parser la date de la session
                session_date = parse_date(session["date"])
                date_str = format_date_fr(session_date)
                
                # Construire l'heure avec le type de session
                heure_str = f"{session['debut']} - {session['fin']}"
                if session.get("type"):
                    heure_str += f" ({session['type']})"
                
                # Lieu de la session ou lieu général
                lieu_str = session.get("lieu", lieu_general)
                
                # Ajouter une nouvelle ligne
                new_row = table.add_row()
                
                # Remplir les cellules avec le bon formatage
                for i, text in enumerate([date_str, heure_str, lieu_str]):
                    cell = new_row.cells[i]
                    cell.text = ""
                    para = cell.paragraphs[0]
                    run = para.add_run(text)
                    # Appliquer la police Calibri 11pt
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
            
            # Appliquer les bordures au tableau
            tbl = table._tbl
            tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            tbl_borders = parse_xml(
                r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'</w:tblBorders>'
            )
            tbl_pr.append(tbl_borders)
            if tbl.tblPr is None:
                tbl.insert(0, tbl_pr)
        
        # Supprimer les paragraphes liés au lien ressources si pas de lien
        if not lien_ressources:
            paragraphs_to_clear = [
                "Vous pourrez vous connecter à cette page",
                "Vous trouverez sur cette page des détails"
            ]
            for para in doc.paragraphs:
                for check in paragraphs_to_clear:
                    if check in para.text:
                        para.clear()
        
        # Nom du fichier
        nom_clean = nom.replace(" ", "_")
        prenom_clean = prenom.replace(" ", "_")
        filename = f"Convocation_{nom_clean}_{prenom_clean}.docx"
        
        # Chemin de sortie
        if output_dir:
            output_path = os.path.join(output_dir, filename)
        elif "_source_dir" in data and data["_source_dir"]:
            output_path = os.path.join(data["_source_dir"], filename)
        else:
            output_path = filename
        
        doc.save(output_path)
        fichiers_generes.append(output_path)
        print(f"   ✅ {filename}")
    
    print(f"\n🎉 {len(fichiers_generes)} convocation(s) générée(s)")
    return fichiers_generes


if __name__ == "__main__":
    if len(sys.argv) > 1:
        json_path = sys.argv[1]
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        source_dir = os.path.dirname(os.path.abspath(json_path))
        if os.path.basename(source_dir) == "data":
            source_dir = os.path.dirname(source_dir)
        if source_dir and source_dir != os.getcwd():
            data["_source_dir"] = source_dir
        generer_convocation(data)
    else:
        print("Usage: python3 generer_convocation.py <fichier.json>")
        print("\nExemple de structure JSON:")
        exemple = {
            "nom_formation": "Intégrer l'IA Générative à votre Activité",
            "date_debut": "25/11/2025",
            "date_fin": "27/11/2025",
            "lieu": "Distanciel",
            "duree_heures": 18,
            "duree_jours": 3,
            "formateurs": ["ALBOUZE Alexis"],
            "apprenants": [
                {"nom": "DUPONT", "prenom": "Jean"}
            ],
            "sessions": [
                {"date": "25/11/2025", "type": "E-learning", "debut": "10h00", "fin": "12h00"},
                {"date": "25/11/2025", "type": "Visio", "debut": "15h00", "fin": "17h00"},
                {"date": "26/11/2025", "type": "E-learning", "debut": "10h00", "fin": "12h00"},
                {"date": "26/11/2025", "type": "Visio", "debut": "15h00", "fin": "17h00"}
            ],
            "date_emission": "20/11/2025",
            "lien_ressources": "https://drive.google.com/..."
        }
        print(json.dumps(exemple, indent=2, ensure_ascii=False))
