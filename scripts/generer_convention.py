#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur de conventions de formation professionnelle pour MINDNESS
Utilise le template "convention template.docx" comme base.
Génère une convention par client/entreprise.
"""

import json
import sys
import os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt
from copy import deepcopy


def format_date_fr(date_obj):
    """Formate une date en français (ex: 16 décembre 2024)."""
    mois = [
        "", "janvier", "février", "mars", "avril", "mai", "juin",
        "juillet", "août", "septembre", "octobre", "novembre", "décembre"
    ]
    return f"{date_obj.day} {mois[date_obj.month]} {date_obj.year}"


def format_date_short(date_obj):
    """Formate une date en format court (ex: 16/12/2024)."""
    return date_obj.strftime("%d/%m/%Y")


def parse_date(date_str):
    """Parse une date en différents formats."""
    formats = ["%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%d/%m/%y"]
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    raise ValueError(f"Format de date non reconnu: {date_str}")


def replace_in_paragraph(para, old_text, new_text):
    """Remplace le texte dans un paragraphe en préservant le formatage si possible."""
    if old_text not in para.text:
        return False
    
    # Essayer de remplacer dans chaque run
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    
    # Si fragmenté entre plusieurs runs, reconstruire le paragraphe
    new_full_text = para.text.replace(old_text, new_text)
    para.clear()
    para.add_run(new_full_text)
    return True


def replace_in_cell(cell, old_text, new_text):
    """Remplace le texte dans une cellule de tableau."""
    for para in cell.paragraphs:
        replace_in_paragraph(para, old_text, new_text)


def generer_convention(data: dict, output_dir: str = None, template_path: str = "convention template.docx"):
    """
    Génère une convention de formation professionnelle.
    
    Args:
        data: Dictionnaire contenant les données de la convention
        output_dir: Dossier de sortie (défaut: dossier courant)
        template_path: Chemin vers le template Word
    
    Returns:
        Chemin du fichier généré
    """
    
    # Trouver le template
    template = Path(template_path)
    if not template.exists():
        script_dir = Path(__file__).parent
        template = script_dir.parent / "templates" / template_path
        if not template.exists():
            raise FileNotFoundError(f"Template non trouvé: {template_path}")
    template_path = str(template)
    
    # Parser les dates
    date_debut = parse_date(data["date_debut"])
    date_fin = parse_date(data["date_fin"])
    
    # Date de signature (par défaut: aujourd'hui)
    if "date_signature" in data and data["date_signature"]:
        date_signature = parse_date(data["date_signature"])
    else:
        date_signature = datetime.now()
    
    # Lieu de signature (par défaut: Paris)
    lieu_signature = data.get("lieu_signature", "Paris")
    
    # Informations bénéficiaire
    beneficiaire = data.get("beneficiaire", {})
    nom_entreprise = beneficiaire.get("nom", "")
    siren = beneficiaire.get("siren", "")
    siret = beneficiaire.get("siret", "")
    adresse_siege = beneficiaire.get("adresse", "")
    representant = beneficiaire.get("representant", "")
    fonction_representant = beneficiaire.get("fonction", "")
    
    # Informations formation
    nom_formation = data.get("nom_formation", "")
    objectif_professionnel = data.get("objectif_professionnel", "")
    contenu_pedagogique = data.get("contenu_pedagogique", [])
    moyens_pedagogiques = data.get("moyens_pedagogiques", "")
    moyens_fournis_beneficiaire = data.get("moyens_fournis_beneficiaire", "")
    
    # Modalité et lieu
    modalite = data.get("modalite", "Présentiel")
    taux_distance = data.get("taux_distance", "0")
    lieu = data.get("lieu", "")
    
    # Durée
    duree_heures = data.get("duree_heures", 0)
    duree_jours = data.get("duree_jours", 0)
    periode_dates = data.get("periode_dates", "")
    horaires = data.get("horaires", "")
    
    # Effectif
    effectif_min = data.get("effectif_min", 2)
    effectif_max = data.get("effectif_max", 12)
    
    # Prix
    prix_ht = data.get("prix_ht", "")
    prix_ttc = data.get("prix_ttc", "")
    frais_deplacement = data.get("frais_deplacement", "")
    acompte = data.get("acompte", "")
    solde = data.get("solde", "")
    
    # Participants
    apprenants = data.get("apprenants", [])
    
    print(f"📝 Génération de la convention de formation")
    print(f"   Bénéficiaire : {nom_entreprise}")
    print(f"   Formation : {nom_formation}")
    print(f"   Du {format_date_fr(date_debut)} au {format_date_fr(date_fin)}")
    print(f"   Durée : {duree_heures} heures ({duree_jours} jours)")
    print(f"   {len(apprenants)} participant(s)")
    
    # Charger le template
    doc = Document(template_path)
    
    # Préparer les remplacements
    replacements = {
        # Bénéficiaire
        "{{NOM_ENTREPRISE}}": nom_entreprise,
        "{{SIREN}}": siren,
        "{{SIRET}}": siret,
        "{{ADRESSE_SIEGE}}": adresse_siege,
        "{{REPRESENTANT}}": representant,
        "{{FONCTION_REPRESENTANT}}": fonction_representant,
        
        # Formation
        "{{NOM_FORMATION}}": nom_formation,
        "{{OBJECTIF_PROFESSIONNEL}}": objectif_professionnel,
        "{{MOYENS_PEDAGOGIQUES}}": moyens_pedagogiques,
        "{{MOYENS_FOURNIS_BENEFICIAIRE}}": moyens_fournis_beneficiaire,
        
        # Modalité et lieu
        "{{MODALITE}}": modalite,
        "{{TAUX_DISTANCE}}": str(taux_distance),
        "{{LIEU}}": lieu,
        
        # Durée
        "{{DUREE_HEURES}}": str(duree_heures),
        "{{DUREE_JOURS}}": str(duree_jours),
        "{{PERIODE_DATES}}": periode_dates,
        "{{HORAIRES}}": horaires,
        
        # Effectif
        "{{EFFECTIF_MIN}}": str(effectif_min),
        "{{EFFECTIF_MAX}}": str(effectif_max),
        
        # Prix
        "{{PRIX_HT}}": prix_ht,
        "{{PRIX_TTC}}": prix_ttc,
        "{{FRAIS_DEPLACEMENT}}": frais_deplacement,
        "{{ACOMPTE}}": acompte,
        "{{SOLDE}}": solde,
        
        # Signature
        "{{DATE_SIGNATURE}}": format_date_short(date_signature),
        "{{LIEU_SIGNATURE}}": lieu_signature,
    }
    
    # Remplacements dans les paragraphes
    for para in doc.paragraphs:
        for old, new in replacements.items():
            replace_in_paragraph(para, old, new)
        
        # Gérer le contenu pédagogique (liste)
        if "{{CONTENU_PEDAGOGIQUE}}" in para.text:
            # Remplacer par la liste des items
            if contenu_pedagogique:
                para.clear()
                for i, item in enumerate(contenu_pedagogique):
                    if i > 0:
                        para.add_run("\n")
                    para.add_run(f"• {item}")
            else:
                replace_in_paragraph(para, "{{CONTENU_PEDAGOGIQUE}}", "")
    
    # Remplacements dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in replacements.items():
                    replace_in_cell(cell, old, new)
    
    # Gérer le tableau des participants (tableau 0 généralement)
    if len(doc.tables) > 0 and apprenants:
        # Trouver le tableau des participants (celui avec Nom | Prénom | Fonction | E-mail)
        for table in doc.tables:
            header_text = " ".join([cell.text for cell in table.rows[0].cells]).lower()
            if "nom" in header_text and "prénom" in header_text:
                # C'est le tableau des participants
                # Supprimer les lignes existantes sauf l'en-tête
                while len(table.rows) > 1:
                    tr = table.rows[-1]._tr
                    table._tbl.remove(tr)
                
                # Ajouter une ligne par participant
                for apprenant in apprenants:
                    new_row = table.add_row()
                    cells_data = [
                        apprenant.get("nom", ""),
                        apprenant.get("prenom", ""),
                        apprenant.get("fonction", ""),
                        apprenant.get("email", "")
                    ]
                    for i, text in enumerate(cells_data):
                        if i < len(new_row.cells):
                            cell = new_row.cells[i]
                            cell.text = ""
                            para = cell.paragraphs[0]
                            run = para.add_run(text)
                            run.font.name = "Calibri"
                            run.font.size = Pt(11)
                break
    
    # Nom du fichier
    nom_clean = nom_entreprise.replace(" ", "_").replace("/", "-")
    filename = f"Convention_{nom_clean}.docx"
    
    # Chemin de sortie
    if output_dir:
        output_path = os.path.join(output_dir, filename)
    elif "_source_dir" in data and data["_source_dir"]:
        output_path = os.path.join(data["_source_dir"], filename)
    else:
        output_path = filename
    
    doc.save(output_path)
    print(f"   ✅ {filename}")
    print(f"\n🎉 Convention générée : {output_path}")
    
    return output_path


if __name__ == "__main__":
    if len(sys.argv) > 1:
        json_path = sys.argv[1]
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        source_dir = os.path.dirname(os.path.abspath(json_path))
        if os.path.basename(source_dir) == "data":
            source_dir = os.path.dirname(source_dir)
        if source_dir and source_dir != os.getcwd():
            data["_source_dir"] = source_dir
        generer_convention(data)
    else:
        print("Usage: python3 generer_convention.py <fichier.json>")
        print("\nExemple de structure JSON:")
        exemple = {
            "beneficiaire": {
                "nom": "ENTREPRISE XYZ",
                "siren": "123456789",
                "siret": "12345678900011",
                "adresse": "123 RUE EXEMPLE 75001 PARIS",
                "representant": "Jean DUPONT",
                "fonction": "Directeur Général"
            },
            "nom_formation": "Intégrer l'IA Générative à votre Activité",
            "objectif_professionnel": "À l'issue de la formation, le stagiaire sait...",
            "contenu_pedagogique": [
                "Maîtriser les concepts fondamentaux de l'IA générative",
                "Rédiger des prompts efficaces",
                "Intégrer l'IA dans les processus métier"
            ],
            "moyens_pedagogiques": "Exposés interactifs, démonstrations, études de cas...",
            "moyens_fournis_beneficiaire": "Les comptes/licences logiciels sont fournis par le Bénéficiaire.",
            "modalite": "Présentiel",
            "taux_distance": "0",
            "lieu": "Locaux du client à Paris",
            "date_debut": "25/11/2025",
            "date_fin": "27/11/2025",
            "duree_heures": 18,
            "duree_jours": 3,
            "periode_dates": "25, 26 et 27 novembre 2025",
            "horaires": "9h00–12h00 et 14h00–17h00",
            "effectif_min": 2,
            "effectif_max": 12,
            "prix_ht": "5 400 € HT",
            "prix_ttc": "6 480 € TTC",
            "frais_deplacement": "500 € HT (600 € TTC)",
            "acompte": "30 % à la signature",
            "solde": "70 % à la fin de la formation",
            "apprenants": [
                {"nom": "MARTIN", "prenom": "Marie", "fonction": "Chef de projet", "email": "marie.martin@xyz.com"},
                {"nom": "DURAND", "prenom": "Pierre", "fonction": "Développeur", "email": "pierre.durand@xyz.com"}
            ],
            "lieu_signature": "Paris",
            "date_signature": "04/11/2025"
        }
        print(json.dumps(exemple, indent=2, ensure_ascii=False))

