#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur de feuilles d'émargement pour formations MINDNESS
Utilise le template EMARGEMENT TEMPLATE.docx comme base.
"""

import json
import sys
import os
import copy
from datetime import datetime, timedelta
from pathlib import Path
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def format_date_fr(date_obj):
    """Formate une date en français (ex: lundi 16 décembre 2024)."""
    mois = [
        "", "janvier", "février", "mars", "avril", "mai", "juin",
        "juillet", "août", "septembre", "octobre", "novembre", "décembre"
    ]
    jours = ["lundi", "mardi", "mercredi", "jeudi", "vendredi", "samedi", "dimanche"]
    return f"{jours[date_obj.weekday()]} {date_obj.day} {mois[date_obj.month]} {date_obj.year}"


def format_date_short(date_obj):
    """Formate une date en format court (ex: 16/12/2024)."""
    return date_obj.strftime("%d/%m/%Y")


def parse_date(date_str):
    """Parse une date en différents formats."""
    formats = ["%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%d/%m/%y"]
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    raise ValueError(f"Format de date non reconnu: {date_str}")


def generate_date_range(date_debut, date_fin):
    """Génère la liste des jours ouvrés entre deux dates."""
    dates = []
    current = date_debut
    while current <= date_fin:
        if current.weekday() < 5:
            dates.append(current)
        current += timedelta(days=1)
    return dates


def copy_element(element):
    """Copie profonde d'un élément XML."""
    return copy.deepcopy(element)


def set_cell_shading(cell, color="D3D3D3"):
    """Applique un fond gris à une cellule."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="000000", size="6"):
    """Applique des bordures à une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        border.set(qn('w:space'), '0')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_table_borders(table):
    """Applique des bordures à tout le tableau."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:color'), '000000')
        border.set(qn('w:space'), '0')
        tblBorders.append(border)
    
    # Supprimer anciennes bordures si elles existent
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)
    
    tblPr.append(tblBorders)
    
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def generer_emargement(data: dict, output_path: str = None, template_path: str = "EMARGEMENT TEMPLATE.docx"):
    """
    Génère une feuille d'émargement complète.
    Regroupe les sessions du même jour sur une seule page.
    """
    
    # Trouver le template
    template = Path(template_path)
    if not template.exists():
        # Chercher dans le dossier templates/ (à côté du dossier scripts/)
        script_dir = Path(__file__).parent
        template = script_dir.parent / "templates" / template_path
        if not template.exists():
            raise FileNotFoundError(f"Template non trouvé: {template_path}")
    template_path = str(template)
    
    date_debut = parse_date(data["date_debut"])
    date_fin = parse_date(data["date_fin"])
    
    print(f"📅 Formation du {format_date_short(date_debut)} au {format_date_short(date_fin)}")
    
    # Charger le template
    doc = Document(template_path)
    
    # Préparer les valeurs communes
    date_debut_str = format_date_short(date_debut)
    date_fin_str = format_date_short(date_fin)
    ville_signature = data.get("ville_signature", "Paris")
    formateurs_list = data["formateurs"]
    
    # Collecter les éléments du template
    body = doc.element.body
    template_elements = []
    sectPr = None
    
    for elem in list(body):
        tag = elem.tag.split('}')[-1]
        
        if tag == 'sectPr':
            sectPr = elem
            continue
        
        if tag == 'p':
            has_text = False
            for t in elem.iter(qn('w:t')):
                if t.text and t.text.strip():
                    has_text = True
                    break
            if not has_text:
                all_text = "".join(t.text or "" for t in elem.iter(qn('w:t')))
                if not all_text.strip():
                    continue
        
        if tag == 'p':
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                pb = pPr.find(qn('w:pageBreakBefore'))
                if pb is not None:
                    pPr.remove(pb)
        
        template_elements.append(elem)
    
    # Vider le body
    for elem in list(body):
        if not elem.tag.endswith('sectPr'):
            body.remove(elem)
    
    # Regrouper les sessions par jour
    if "sessions" in data:
        sessions_by_day = defaultdict(list)
        for session in data["sessions"]:
            session_date = parse_date(session["date"])
            sessions_by_day[session_date].append({
                "type": session.get("type", ""),
                "horaires": f"{session['debut']}-{session['fin']}",
                "debut": session["debut"],
                "fin": session["fin"]
            })
        
        # Trier les jours
        jours = sorted(sessions_by_day.keys())
        pages_to_generate = [(jour, sessions_by_day[jour]) for jour in jours]
        print(f"📆 {len(pages_to_generate)} jour(s) de formation")
    else:
        # Format classique
        jours_formation = generate_date_range(date_debut, date_fin)
        horaires = data["horaires"]
        pages_to_generate = []
        for jour in jours_formation:
            sessions = [
                {"type": "Matin", "horaires": f"{horaires['matin']['debut']}-{horaires['matin']['fin']}"},
                {"type": "Après-midi", "horaires": f"{horaires['apres_midi']['debut']}-{horaires['apres_midi']['fin']}"}
            ]
            pages_to_generate.append((jour, sessions))
        print(f"📆 {len(pages_to_generate)} jour(s) de formation")
    
    # Générer les pages
    for page_idx, (jour, sessions) in enumerate(pages_to_generate):
        date_jour_str = format_date_fr(jour)
        
        sessions_str = ", ".join([f"{s['type']} {s['horaires']}" for s in sessions])
        print(f"   → {date_jour_str} : {sessions_str}")
        
        # Copier les éléments du template
        page_elements = [copy_element(elem) for elem in template_elements]
        
        # Ajouter pageBreakBefore pour les pages suivantes
        for i, elem in enumerate(page_elements):
            if page_idx > 0 and i == 0:
                pPr = elem.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    elem.insert(0, pPr)
                pageBreakBefore = OxmlElement('w:pageBreakBefore')
                pageBreakBefore.set(qn('w:val'), '1')
                pPr.append(pageBreakBefore)
            body.append(elem)
            
            if i == 0:
                for _ in range(2):
                    empty_p = OxmlElement('w:p')
                    body.append(empty_p)
    
    if sectPr is not None:
        body.append(sectPr)
    
    # Recharger le document
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    doc = Document(buffer)
    
    # Remplacements dans chaque page
    tables_per_page = 2
    
    for page_idx, (jour, sessions) in enumerate(pages_to_generate):
        jour_str_key = jour.strftime("%Y-%m-%d")
        intervenants_jour = data.get("intervenants_par_jour", {}).get(jour_str_key, data["formateurs"])
        
        date_jour_str = format_date_fr(jour)
        date_signature = format_date_short(jour)
        
        table_info_idx = page_idx * tables_per_page
        table_emarg_idx = page_idx * tables_per_page + 1
        
        if table_info_idx >= len(doc.tables) or table_emarg_idx >= len(doc.tables):
            continue
        
        table_info = doc.tables[table_info_idx]
        table_emarg = doc.tables[table_emarg_idx]
        
        # TABLEAU 1 : INFOS FORMATION
        # Compter les occurrences de DATE pour savoir laquelle remplacer
        date_count = [0]  # Utiliser une liste pour modifier dans la closure
        
        for row in table_info.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text
                        text = text.replace("XXXXX", data["nom_formation"])
                        text = text.replace("PRESENTIELOUDISTANCIEL, LIEU", data["lieu"])
                        text = text.replace("NOMBREHEURES", str(data["duree_heures"]))
                        # Remplacer DATE selon l'ordre d'apparition
                        while "DATE" in text:
                            if date_count[0] == 0:
                                text = text.replace("DATE", date_debut_str, 1)
                            else:
                                text = text.replace("DATE", date_fin_str, 1)
                            date_count[0] += 1
                        text = text.replace("M. ALBOUZE Alexis", formateurs_list[0])
                        text = text.replace("ALBOUZE Alexis", formateurs_list[0])
                        run.text = text
        
        # TABLEAU 2 : ÉMARGEMENT - Reconstruire complètement
        # Appliquer les bordures au tableau
        set_table_borders(table_emarg)
        
        # Ligne 0 : Date du jour
        for cell in table_emarg.rows[0].cells:
            set_cell_borders(cell)
            for para in cell.paragraphs:
                para.clear()
                run = para.add_run(f"Date : {date_jour_str}")
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
                run.bold = True
        
        # Supprimer toutes les lignes existantes sauf la première (en-tête date)
        while len(table_emarg.rows) > 1:
            tr = table_emarg.rows[-1]._tr
            table_emarg._tbl.remove(tr)
        
        # Ajouter les créneaux
        for session in sessions:
            # Ligne créneau (fusionnée, fond gris)
            row_creneau = table_emarg.add_row()
            cell_creneau = row_creneau.cells[0]
            cell_creneau.merge(row_creneau.cells[1])
            set_cell_shading(cell_creneau)
            set_cell_borders(cell_creneau)
            para = cell_creneau.paragraphs[0]
            para.clear()
            run = para.add_run(f"Créneau : {session['horaires']} ({session['type']})")
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.bold = True
            
            # Lignes apprenants
            for apprenant in data["apprenants"]:
                row_app = table_emarg.add_row()
                cell_nom = row_app.cells[0]
                cell_sig = row_app.cells[1]
                set_cell_borders(cell_nom)
                set_cell_borders(cell_sig)
                para = cell_nom.paragraphs[0]
                para.clear()
                run = para.add_run(f"{apprenant['nom']} {apprenant['prenom']}  ---  {apprenant['email']}")
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
                cell_sig.paragraphs[0].clear()
                
                tr = row_app._tr
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), "500")
                trHeight.set(qn('w:hRule'), "atLeast")
                trPr.append(trHeight)
            
            # Ligne "Formateur" (fond gris)
            row_form_label = table_emarg.add_row()
            cell_form_label = row_form_label.cells[0]
            cell_form_label.merge(row_form_label.cells[1])
            set_cell_shading(cell_form_label)
            set_cell_borders(cell_form_label)
            para = cell_form_label.paragraphs[0]
            para.clear()
            run = para.add_run("Formateur")
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.bold = True
            
            # Lignes intervenants
            for intervenant in intervenants_jour:
                row_int = table_emarg.add_row()
                cell_nom = row_int.cells[0]
                cell_sig = row_int.cells[1]
                set_cell_borders(cell_nom)
                set_cell_borders(cell_sig)
                para = cell_nom.paragraphs[0]
                para.clear()
                run = para.add_run(intervenant)
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
                cell_sig.paragraphs[0].clear()
                
                tr = row_int._tr
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), "500")
                trHeight.set(qn('w:hRule'), "atLeast")
                trPr.append(trHeight)
    
    # Remplacer "Fait à"
    page_idx = 0
    for para in doc.paragraphs:
        text = para.text
        if "Fait" in text and "xx" in text:
            if page_idx < len(pages_to_generate):
                jour, _ = pages_to_generate[page_idx]
                date_sig = format_date_short(jour)
                para.clear()
                run = para.add_run(f"Fait à {ville_signature}, le {date_sig}")
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
                page_idx += 1
    
    # Sauvegarde (dans le même dossier que le fichier JSON source si spécifié)
    if output_path is None:
        nom_clean = "".join(c if c.isalnum() or c in " -_" else "" for c in data["nom_formation"])
        nom_clean = nom_clean.replace(" ", "_")[:50]
        filename = f"Emargement_{nom_clean}_{date_debut.strftime('%Y%m%d')}.docx"
        
        # Si un dossier source est spécifié dans data, l'utiliser
        if "_source_dir" in data and data["_source_dir"]:
            output_path = os.path.join(data["_source_dir"], filename)
        else:
            output_path = filename
    
    doc.save(output_path)
    print(f"✅ Feuille d'émargement générée : {output_path}")
    return output_path


if __name__ == "__main__":
    if len(sys.argv) > 1:
        json_path = sys.argv[1]
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        # Stocker le dossier client pour y générer le fichier
        # Si le JSON est dans un sous-dossier data/, remonter au dossier client
        source_dir = os.path.dirname(os.path.abspath(json_path))
        if os.path.basename(source_dir) == "data":
            source_dir = os.path.dirname(source_dir)  # Remonter au dossier client
        if source_dir and source_dir != os.getcwd():
            data["_source_dir"] = source_dir
        generer_emargement(data)
    else:
        exemple_data = {
            "nom_formation": "Prompt Engineering Avancé",
            "date_debut": "16/12/2024",
            "date_fin": "17/12/2024",
            "lieu": "Présentiel, 12 rue de l'Innovation, Paris",
            "duree_heures": 14,
            "horaires": {
                "matin": {"debut": "9h30", "fin": "12h30"},
                "apres_midi": {"debut": "14h00", "fin": "17h30"}
            },
            "formateurs": ["ALBOUZE Alexis"],
            "apprenants": [
                {"nom": "DUPONT", "prenom": "Jean", "email": "jean.dupont@example.com"},
                {"nom": "MARTIN", "prenom": "Marie", "email": "marie.martin@example.com"},
                {"nom": "BERNARD", "prenom": "Pierre", "email": "pierre.bernard@example.com"},
            ],
            "ville_signature": "Paris"
        }
        
        print("📋 Génération d'exemple...")
        generer_emargement(exemple_data)
