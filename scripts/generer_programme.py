#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur de programmes pédagogiques pour formations MINDNESS
Utilise le template "programme_pedagogique template.docx" comme base.
Génère un programme pédagogique par formation.
"""

import json
import sys
import os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def format_date_fr(date_obj):
    """Formate une date en français (ex: 16 décembre 2024)."""
    mois = [
        "", "janvier", "février", "mars", "avril", "mai", "juin",
        "juillet", "août", "septembre", "octobre", "novembre", "décembre"
    ]
    return f"{date_obj.day} {mois[date_obj.month]} {date_obj.year}"


def parse_date(date_str):
    """Parse une date en différents formats."""
    formats = ["%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%d/%m/%y"]
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    raise ValueError(f"Format de date non reconnu: {date_str}")


def replace_in_paragraph(para, old_text, new_text):
    """Remplace le texte dans un paragraphe en préservant le formatage si possible."""
    if old_text not in para.text:
        return False
    
    # Essayer de remplacer dans chaque run
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    
    # Si fragmenté entre plusieurs runs, reconstruire le paragraphe
    new_full_text = para.text.replace(old_text, new_text)
    para.clear()
    para.add_run(new_full_text)
    return True


def replace_in_cell(cell, old_text, new_text):
    """Remplace le texte dans une cellule de tableau."""
    for para in cell.paragraphs:
        replace_in_paragraph(para, old_text, new_text)


def format_list_to_bullets(items):
    """Formate une liste en texte avec puces."""
    if not items:
        return ""
    if isinstance(items, str):
        return items
    return "\n".join([f"• {item}" for item in items])


def generer_programme(data: dict, output_dir: str = None, template_path: str = "programme_pedagogique template.docx"):
    """
    Génère un programme pédagogique de formation.
    
    Args:
        data: Dictionnaire contenant les données de la formation
        output_dir: Dossier de sortie (défaut: dossier courant)
        template_path: Chemin vers le template Word
    
    Returns:
        Chemin du fichier généré
    """
    
    # Trouver le template
    template = Path(template_path)
    if not template.exists():
        script_dir = Path(__file__).parent
        template = script_dir.parent / "templates" / template_path
        if not template.exists():
            raise FileNotFoundError(f"Template non trouvé: {template_path}")
    template_path = str(template)
    
    # Extraire les données
    nom_formation = data.get("nom_formation", "")
    duree_heures = data.get("duree_heures", 0)
    duree_jours = data.get("duree_jours", 0)
    modalite = data.get("modalite", data.get("lieu", ""))
    lieu = data.get("lieu", "")
    
    # Formateur(s)
    formateurs = data.get("formateurs", [])
    if isinstance(formateurs, list):
        formateur_str = ", ".join(formateurs)
    else:
        formateur_str = formateurs
    
    # Données pédagogiques spécifiques au programme
    public_vise = data.get("public_vise", "Tout public")
    prerequis = data.get("prerequis", "Aucun prérequis spécifique")
    accessibilite = data.get("accessibilite", "Formation accessible aux personnes en situation de handicap. Contactez-nous pour adapter les modalités.")
    
    # Objectifs pédagogiques (liste)
    objectifs = data.get("objectifs_pedagogiques", [])
    objectifs_str = format_list_to_bullets(objectifs)
    
    # Modules/Programme détaillé
    modules = data.get("modules", [])
    
    # Méthodes pédagogiques
    methodes = data.get("methodes_pedagogiques", [
        "Exposés interactifs et démonstrations",
        "Exercices pratiques et mises en situation",
        "Études de cas concrets",
        "Échanges et partages d'expériences"
    ])
    methodes_str = format_list_to_bullets(methodes)
    
    # Modalités d'évaluation
    evaluations = data.get("modalites_evaluation", [
        "Évaluation diagnostique en début de formation",
        "Évaluations formatives tout au long de la formation",
        "Évaluation sommative en fin de formation",
        "Questionnaire de satisfaction"
    ])
    evaluations_str = format_list_to_bullets(evaluations)
    
    # Sanction de la formation
    sanction = data.get("sanction_formation", "Attestation de fin de formation et certificat de réalisation remis au stagiaire.")
    
    print(f"📚 Génération du programme pédagogique")
    print(f"   Formation : {nom_formation}")
    print(f"   Durée : {duree_heures} heures ({duree_jours} jours)")
    print(f"   Modalité : {modalite or lieu}")
    print(f"   {len(modules)} module(s)")
    
    # Charger le template
    doc = Document(template_path)
    
    # Préparer les remplacements de base
    replacements = {
        "{{NOM_FORMATION}}": nom_formation,
        "{{DUREE_HEURES}}": str(duree_heures),
        "{{DUREE_JOURS}}": str(duree_jours),
        "{{MODALITE}}": modalite or lieu,
        "{{LIEU}}": lieu,
        "{{FORMATEURS}}": formateur_str,
        "{{PUBLIC_VISE}}": public_vise,
        "{{PREREQUIS}}": prerequis,
        "{{ACCESSIBILITE}}": accessibilite,
        "{{OBJECTIFS_PEDAGOGIQUES}}": objectifs_str,
        "{{METHODES_PEDAGOGIQUES}}": methodes_str,
        "{{MODALITES_EVALUATION}}": evaluations_str,
        "{{SANCTION_FORMATION}}": sanction,
    }
    
    # Remplacements dans les paragraphes
    for para in doc.paragraphs:
        # Traitement spécial pour le programme détaillé
        if "{{PROGRAMME_DETAILLE}}" in para.text:
            para.clear()
            if modules:
                for i, module in enumerate(modules):
                    titre = module.get("titre", f"Module {i+1}")
                    duree_module = module.get("duree", "")
                    contenu = module.get("contenu", [])
                    objectifs_module = module.get("objectifs", [])
                    
                    # Titre du module
                    run_titre = para.add_run(f"\n{titre}")
                    run_titre.bold = True
                    run_titre.font.size = Pt(12)
                    run_titre.font.name = "Calibri"
                    
                    if duree_module:
                        run_duree = para.add_run(f" ({duree_module})")
                        run_duree.font.size = Pt(11)
                        run_duree.font.name = "Calibri"
                    
                    # Objectifs du module
                    if objectifs_module:
                        para.add_run("\n")
                        run_obj_title = para.add_run("Objectifs : ")
                        run_obj_title.italic = True
                        run_obj_title.font.size = Pt(11)
                        run_obj_title.font.name = "Calibri"
                        
                        for obj in objectifs_module:
                            run_obj = para.add_run(f"\n  • {obj}")
                            run_obj.font.size = Pt(11)
                            run_obj.font.name = "Calibri"
                    
                    # Contenu du module
                    if contenu:
                        para.add_run("\n")
                        run_cont_title = para.add_run("Contenu : ")
                        run_cont_title.italic = True
                        run_cont_title.font.size = Pt(11)
                        run_cont_title.font.name = "Calibri"
                        
                        for item in contenu:
                            run_item = para.add_run(f"\n  • {item}")
                            run_item.font.size = Pt(11)
                            run_item.font.name = "Calibri"
                    
                    para.add_run("\n")
            else:
                run = para.add_run("Programme détaillé à définir.")
                run.font.name = "Calibri"
                run.font.size = Pt(11)
            continue
        
        # Remplacements standards
        for old, new in replacements.items():
            replace_in_paragraph(para, old, new)
    
    # Remplacements dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in replacements.items():
                    replace_in_cell(cell, old, new)
    
    # Nom du fichier
    nom_clean = nom_formation[:50].replace(" ", "_").replace("/", "-").replace("'", "")
    filename = f"Programme_pedagogique_{nom_clean}.docx"
    
    # Chemin de sortie
    if output_dir:
        output_path = os.path.join(output_dir, filename)
    elif "_source_dir" in data and data["_source_dir"]:
        output_path = os.path.join(data["_source_dir"], filename)
    else:
        output_path = filename
    
    doc.save(output_path)
    print(f"   ✅ {filename}")
    print(f"\n🎉 Programme pédagogique généré : {output_path}")
    
    return output_path


if __name__ == "__main__":
    if len(sys.argv) > 1:
        json_path = sys.argv[1]
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        source_dir = os.path.dirname(os.path.abspath(json_path))
        if os.path.basename(source_dir) == "data":
            source_dir = os.path.dirname(source_dir)
        if source_dir and source_dir != os.getcwd():
            data["_source_dir"] = source_dir
        generer_programme(data)
    else:
        print("Usage: python3 generer_programme.py <fichier.json>")
        print("\nExemple de structure JSON:")
        exemple = {
            "nom_formation": "Intégrer l'IA Générative à votre Activité",
            "duree_heures": 18,
            "duree_jours": 3,
            "modalite": "Présentiel",
            "lieu": "Paris",
            "formateurs": ["ALBOUZE Alexis"],
            "public_vise": "Professionnels souhaitant intégrer l'IA dans leur activité",
            "prerequis": "Maîtrise des outils bureautiques de base",
            "accessibilite": "Formation accessible aux personnes en situation de handicap. Contactez-nous pour adapter les modalités.",
            "objectifs_pedagogiques": [
                "Comprendre les fondamentaux de l'IA générative",
                "Maîtriser les techniques de prompt engineering",
                "Savoir intégrer l'IA dans les processus métier",
                "Utiliser l'IA de manière éthique et responsable"
            ],
            "modules": [
                {
                    "titre": "Module 1 : Introduction à l'IA Générative",
                    "duree": "3 heures",
                    "objectifs": [
                        "Comprendre ce qu'est l'IA générative",
                        "Connaître les principaux outils disponibles"
                    ],
                    "contenu": [
                        "Histoire et évolution de l'IA",
                        "Différence entre IA traditionnelle et IA générative",
                        "Panorama des outils : ChatGPT, Claude, Midjourney, etc.",
                        "Cas d'usage en entreprise"
                    ]
                },
                {
                    "titre": "Module 2 : Prompt Engineering",
                    "duree": "6 heures",
                    "objectifs": [
                        "Rédiger des prompts efficaces",
                        "Optimiser les résultats obtenus"
                    ],
                    "contenu": [
                        "Anatomie d'un bon prompt",
                        "Techniques avancées : few-shot, chain-of-thought",
                        "Exercices pratiques",
                        "Correction et amélioration des résultats"
                    ]
                },
                {
                    "titre": "Module 3 : Intégration dans les Processus Métier",
                    "duree": "6 heures",
                    "objectifs": [
                        "Identifier les opportunités d'intégration de l'IA",
                        "Mettre en œuvre des solutions concrètes"
                    ],
                    "contenu": [
                        "Audit des processus existants",
                        "Identification des tâches automatisables",
                        "Mise en place de workflows intégrant l'IA",
                        "Mesure du ROI"
                    ]
                },
                {
                    "titre": "Module 4 : Éthique et Bonnes Pratiques",
                    "duree": "3 heures",
                    "objectifs": [
                        "Utiliser l'IA de manière responsable",
                        "Connaître le cadre réglementaire"
                    ],
                    "contenu": [
                        "Biais et limites de l'IA",
                        "Protection des données personnelles (RGPD)",
                        "Propriété intellectuelle",
                        "Charte d'utilisation en entreprise"
                    ]
                }
            ],
            "methodes_pedagogiques": [
                "Exposés interactifs avec démonstrations en direct",
                "Exercices pratiques sur les outils d'IA",
                "Études de cas concrets",
                "Travaux de groupe et partages d'expériences"
            ],
            "modalites_evaluation": [
                "Évaluation diagnostique en début de formation (quiz)",
                "Évaluations formatives (exercices pratiques notés)",
                "Évaluation sommative en fin de formation (projet)",
                "Questionnaire de satisfaction à chaud"
            ],
            "sanction_formation": "Attestation de fin de formation mentionnant les objectifs, la durée et les résultats de l'évaluation. Certificat de réalisation remis au stagiaire."
        }
        print(json.dumps(exemple, indent=2, ensure_ascii=False))
