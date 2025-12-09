#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur de certificats de réalisation pour formations MINDNESS
Utilise le template "certificat_de_réalisation template.docx" comme base.
Génère un certificat par apprenant.
"""

import json
import sys
import os
from datetime import datetime
from pathlib import Path
from docx import Document


def format_date_fr(date_obj):
    """Formate une date en français (ex: 16 décembre 2024)."""
    mois = [
        "", "janvier", "février", "mars", "avril", "mai", "juin",
        "juillet", "août", "septembre", "octobre", "novembre", "décembre"
    ]
    return f"{date_obj.day} {mois[date_obj.month]} {date_obj.year}"


def format_date_short(date_obj):
    """Formate une date en format court (ex: 16/12/2024)."""
    return date_obj.strftime("%d/%m/%Y")


def parse_date(date_str):
    """Parse une date en différents formats."""
    formats = ["%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%d/%m/%y"]
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    raise ValueError(f"Format de date non reconnu: {date_str}")


def generer_certificat(data: dict, output_dir: str = None, template_path: str = "certificat_de_réalisation template.docx"):
    """
    Génère un certificat de réalisation par apprenant.
    
    Args:
        data: Dictionnaire contenant les données de la formation
        output_dir: Dossier de sortie (défaut: dossier courant)
        template_path: Chemin vers le template Word
    
    Returns:
        Liste des chemins des fichiers générés
    """
    
    # Trouver le template
    template = Path(template_path)
    if not template.exists():
        # Chercher dans le dossier templates/ (à côté du dossier scripts/)
        script_dir = Path(__file__).parent
        template = script_dir.parent / "templates" / template_path
        if not template.exists():
            raise FileNotFoundError(f"Template non trouvé: {template_path}")
    template_path = str(template)
    
    # Parser les dates
    date_debut = parse_date(data["date_debut"])
    date_fin = parse_date(data["date_fin"])
    
    # Date de signature : personnalisée ou date de fin par défaut
    if "date_signature" in data and data["date_signature"]:
        date_signature = parse_date(data["date_signature"])
    else:
        date_signature = date_fin
    
    lieu_signature = data.get("lieu_signature", "Paris")
    
    print(f"📜 Génération des certificats de réalisation")
    print(f"   Formation : {data['nom_formation']}")
    print(f"   Du {format_date_short(date_debut)} au {format_date_short(date_fin)}")
    print(f"   Durée : {data['duree_heures']} heures")
    print(f"   {len(data['apprenants'])} apprenant(s)")
    
    fichiers_generes = []
    
    for apprenant in data["apprenants"]:
        nom = apprenant["nom"]
        prenom = apprenant["prenom"]
        nom_complet = f"{nom} {prenom}"
        
        # Charger le template
        doc = Document(template_path)
        
        # Remplacements dans les paragraphes
        for para in doc.paragraphs:
            # D'abord tenter les remplacements run par run
            for run in para.runs:
                text = run.text
                
                # Données de l'apprenant et formation
                text = text.replace("NOM PRENOM", nom_complet)
                text = text.replace("NOM DE LA FORMATION", data["nom_formation"])
                text = text.replace("DATE DÉBUT", format_date_short(date_debut))
                text = text.replace("DATE FIN", format_date_short(date_fin))
                text = text.replace("NOMBREHEURES", str(data["duree_heures"]))
                
                run.text = text
            
            # Gérer les cas où le texte est fragmenté entre plusieurs runs
            # (ex: "Fait à : DATE" peut être sur plusieurs runs)
            full_text = para.text
            
            # Correction de l'inversion dans le template
            if "Fait à : DATE" in full_text:
                # Reconstruire le paragraphe avec le bon texte
                new_text = full_text.replace("Fait à : DATE", f"Fait à : {lieu_signature}")
                para.clear()
                para.add_run(new_text)
            elif "LIEU" in full_text and "Le" in full_text:
                # Gérer l'espace insécable (\xa0) dans le template
                new_text = f"Le : {format_date_short(date_signature)}"
                para.clear()
                para.add_run(new_text)
        
        # Remplacements dans les tableaux (au cas où)
        # ATTENTION: ne pas toucher aux runs qui contiennent des images
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            # Vérifier si le run contient une image
                            from docx.oxml.ns import qn
                            if run._element.findall('.//' + qn('w:drawing')):
                                continue  # Ne pas toucher aux images
                            
                            original_text = run.text
                            text = original_text
                            text = text.replace("NOM PRENOM", nom_complet)
                            text = text.replace("NOM DE LA FORMATION", data["nom_formation"])
                            text = text.replace("DATE DÉBUT", format_date_short(date_debut))
                            text = text.replace("DATE FIN", format_date_short(date_fin))
                            text = text.replace("NOMBREHEURES", str(data["duree_heures"]))
                            if "Fait à : DATE" in text:
                                text = text.replace("Fait à : DATE", f"Fait à : {lieu_signature}")
                            if "Le : LIEU" in text:
                                text = text.replace("Le : LIEU", f"Le : {format_date_short(date_signature)}")
                            
                            # Ne modifier que si le texte a changé
                            if text != original_text:
                                run.text = text
        
        # Supprimer les paragraphes vides à la fin pour tenir sur une page
        while doc.paragraphs and not doc.paragraphs[-1].text.strip():
            p = doc.paragraphs[-1]._element
            p.getparent().remove(p)
        
        # Nom du fichier : Certificat_NOM_Prenom.docx
        nom_clean = nom.replace(" ", "_")
        prenom_clean = prenom.replace(" ", "_")
        filename = f"Certificat_{nom_clean}_{prenom_clean}.docx"
        
        # Déterminer le chemin de sortie
        if output_dir:
            output_path = os.path.join(output_dir, filename)
        elif "_source_dir" in data and data["_source_dir"]:
            output_path = os.path.join(data["_source_dir"], filename)
        else:
            output_path = filename
        
        doc.save(output_path)
        fichiers_generes.append(output_path)
        print(f"   ✅ {filename}")
    
    print(f"\n🎉 {len(fichiers_generes)} certificat(s) généré(s)")
    return fichiers_generes


if __name__ == "__main__":
    if len(sys.argv) > 1:
        json_path = sys.argv[1]
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        # Stocker le dossier client pour y générer les fichiers
        # Si le JSON est dans un sous-dossier data/, remonter au dossier client
        source_dir = os.path.dirname(os.path.abspath(json_path))
        if os.path.basename(source_dir) == "data":
            source_dir = os.path.dirname(source_dir)  # Remonter au dossier client
        if source_dir and source_dir != os.getcwd():
            data["_source_dir"] = source_dir
        generer_certificat(data)
    else:
        # Exemple d'utilisation
        print("Usage: python3 generer_certificat.py <fichier.json>")
        print("\nExemple de structure JSON:")
        exemple = {
            "nom_formation": "Prompt Engineering Avancé",
            "date_debut": "16/12/2024",
            "date_fin": "17/12/2024",
            "duree_heures": 14,
            "apprenants": [
                {"nom": "DUPONT", "prenom": "Jean"},
                {"nom": "MARTIN", "prenom": "Marie"}
            ],
            "lieu_signature": "Paris",
            "date_signature": "17/12/2024"
        }
        print(json.dumps(exemple, indent=2, ensure_ascii=False))

